"""
Core PII Redaction Engine (redactor.py)

This module provides a production-ready, hybrid PII redaction engine combining
Presidio Analyzer, spaCy NER, custom deterministic regex recognizers, Luhn algorithm validation,
a thread-safe standardized PseudonymMapper, and docx layout-preserving redactor.
"""

import io
import re
import threading
from dataclasses import dataclass
from typing import Union, BinaryIO, List, Dict, Tuple, Any, Optional

import docx
from docx.text.paragraph import Paragraph

from presidio_analyzer import (
    AnalyzerEngine,
    PatternRecognizer,
    Pattern,
    RecognizerResult,
    RecognizerRegistry,
)
from presidio_analyzer.nlp_engine import NlpEngineProvider


# ---------------------------------------------------------------------------
# Legal & Prospectus Vocabulary Exclusion Set (Never-Redact List)
# ---------------------------------------------------------------------------
NON_PII_TERMS = {
    "companies act",
    "sebi",
    "red herring prospectus",
    "book built offer",
    "registered office",
    "corporate office",
    "contact person",
    "website",
    "fresh issue",
    "offer for sale",
    "equity shares",
    "price band",
    "floor price",
    "cap price",
    "risk factors",
    "draft red herring prospectus",
    "bse",
    "nse",
    "icdr regulations",
    "sebi icdr regulations",
    "promoter selling shareholder",
    "promoter selling shareholders",
    "board of directors",
    "key managerial personnel",
    "senior management",
    "audit committee",
    "articles of association",
    "memorandum of association",
    # Section & Table Description Terms
    "details of the offer",
    "details of the offer to public",
    "size of the offer for sale",
    "objects of the offer",
    "basis for offer price",
    "terms of the offer",
    "forward-looking statements",
    "currency of presentation",
    "market data and currency",
    "central government",
    "state government",
    "government of india",
    "indian government",
    "weighted average cost",
    "cost of acquisition",
    "share reservation",
    "foreign ownership",
    "indian securities",
}

NON_PII_WORDS = {
    "company", "companies", "offer", "issue", "equity", "share", "shares",
    "board", "directors", "director", "promoter", "promoters", "shareholder",
    "shareholders", "section", "act", "rules", "regulations", "prospectus",
    "registered", "office", "corporate", "contact", "person", "website",
    "email", "telephone", "phone", "details", "total", "type", "size",
    "fresh", "sale", "price", "band", "floor", "cap", "risk", "factors",
    "general", "bse", "nse", "sebi", "icdr", "scrr", "india", "maharashtra",
    "pune", "mumbai", "draft", "red", "herring", "book", "built", "holding",
    "capital", "value", "face", "amount", "aggregate", "million", "billion",
    "table", "statement", "financial", "rights", "securities", "exchange",
    "bank", "trust", "group", "summary", "schedule", "annexure", "index",
    "definitions", "abbreviations", "conventions", "presentation", "name",
    "key", "officers", "officer", "statutory", "auditors", "auditor",
    "independent", "compliance", "managers", "manager", "registrar",
    "lead", "running", "grievance", "investor", "advisors", "advisor",
    # Additional Section/Header & Geographic Nouns
    "public", "looking", "statements", "market", "currency", "data",
    "government", "central", "state", "taluka", "district", "village",
    "khed", "chakan", "tehsil", "mandal", "acquisition", "weighted",
    "average", "cost", "reservation", "among", "ownership", "foreign",
    "restrictions", "objects", "basis", "benefits", "legal", "other",
    "information", "regulatory", "disclosures", "business", "our",
    "and", "the", "for", "to", "of", "in", "on", "at", "by", "with",
    "from", "into", "or", "an", "a", "is", "are", "was", "were",
}

ALL_CAPS_STOP_WORDS = {
    "LIMITED", "PRIVATE", "PROSPECTUS", "HERRING", "RED", "OFFER",
    "CORPORATE", "REGISTERED", "OFFICE", "DIRECTOR", "SECRETARY",
    "TABLE", "SECTION", "STATEMENT", "FINANCIAL", "RIGHTS", "ISSUE",
    "DRAFT", "ACT", "INDIA", "SECURITIES", "EXCHANGE", "BOARD",
    "BANK", "COMPANY", "TRUST", "GROUP", "TOTAL", "DETAILS", "SUMMARY",
    "DEFINITIONS", "ABBREVIATIONS", "CONVENTIONS", "PRESENTATION",
    "KEY", "OFFICERS", "STATUTORY", "AUDITORS", "INDEPENDENT",
    "COMPLIANCE", "MANAGERS", "REGISTRAR", "LEAD", "RUNNING",
    # Section Headers, Prepositions, Administrative Terms
    "PUBLIC", "LOOKING", "STATEMENTS", "MARKET", "CURRENCY", "DATA",
    "GOVERNMENT", "CENTRAL", "STATE", "TALUKA", "DISTRICT", "VILLAGE",
    "KHED", "CHAKAN", "ACQUISITION", "WEIGHTED", "AVERAGE", "COST",
    "RESERVATION", "AMONG", "OWNERSHIP", "FOREIGN", "RESTRICTIONS",
    "OBJECTS", "BASIS", "BENEFITS", "LEGAL", "OTHER", "INFORMATION",
    "REGULATORY", "DISCLOSURES", "BUSINESS", "OUR", "QIBS", "NIIS",
    "RIIS", "AND", "THE", "FOR", "TO", "OF", "IN", "ON", "AT", "BY",
    "WITH", "FROM", "INTO", "OR", "AN", "A",
}

HEADING_BLOCKLIST = {
    "general information", "capital structure", "definitions", "abbreviations",
    "currency", "presentation", "risk factors", "summary", "introduction",
    "the offer", "summary financial statements", "general risks", "listing",
    "book running lead managers", "registrar to the offer", "bid/offer period",
    "details of the offer", "type", "size of the fresh issue", "total offer size"
}
def is_valid_luhn(card_number_str: str) -> bool:
    """Validate numeric string against the Luhn algorithm checksum."""
    digits = [int(c) for c in card_number_str if c.isdigit()]
    if len(digits) < 13 or len(digits) > 19:
        return False
    checksum = 0
    reverse_digits = digits[::-1]
    for i, d in enumerate(reverse_digits):
        if i % 2 == 1:
            d = d * 2
            if d > 9:
                d -= 9
        checksum += d
    return checksum % 10 == 0


# ---------------------------------------------------------------------------
# Custom Presidio Pattern Recognizers
# ---------------------------------------------------------------------------
class CustomEmailRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="email_regex",
                regex=r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
                score=1.0,
            )
        ]
        super().__init__(
            supported_entity="EMAIL",
            patterns=patterns,
            name="CustomEmailRecognizer",
        )


class CustomPhoneRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="phone_formatted_indian",
                regex=r"\+?\s*91[\s-]?[6-9]\d{4}[\s-]?\d{5}|\b[6-9]\d{4}[\s-]?\d{5}\b|\+?\s*91[\s-]?\d{2,4}[\s-]?\d{3,5}[\s-]?\d{3,5}",
                score=0.95,
            ),
            Pattern(
                name="phone_formatted_intl",
                regex=r"\+\d{1,4}[\s-]?\(?\d{2,4}\)?[\s-]?\d{6,10}",
                score=0.95,
            ),
        ]
        super().__init__(
            supported_entity="PHONE",
            patterns=patterns,
            name="CustomPhoneRecognizer",
        )

    def validate_result(self, pattern_text: str) -> bool:
        """Ensure candidate contains valid phone digit counts and is not a section/page reference."""
        clean = pattern_text.strip()
        digits = [c for c in clean if c.isdigit()]
        if len(digits) < 10 or len(digits) > 15:
            return False
        if len(set(digits)) == 1:
            return False
        return True


class CustomIPAddressRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="ip_address_regex",
                regex=r"\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b",
                score=1.0,
            )
        ]
        super().__init__(
            supported_entity="IP_ADDRESS",
            patterns=patterns,
            name="CustomIPAddressRecognizer",
        )


class CustomCreditCardRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="credit_card_regex",
                regex=r"\b(?:\d[ -]*?){13,16}\b",
                score=1.0,
            )
        ]
        super().__init__(
            supported_entity="CREDIT_CARD",
            patterns=patterns,
            name="CustomCreditCardRecognizer",
        )

    def validate_result(self, pattern_text: str) -> bool:
        """Validate candidate credit card match using Luhn checksum."""
        return is_valid_luhn(pattern_text)


class CustomGovtIDRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="us_ssn",
                regex=r"\b\d{3}-\d{2}-\d{4}\b",
                score=1.0,
            ),
            Pattern(
                name="indian_pan",
                regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b",
                score=1.0,
            ),
            Pattern(
                name="indian_aadhaar",
                regex=r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b",
                score=0.95,
            ),
            Pattern(
                name="corporate_identity_number_cin",
                regex=r"\b[UL]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b",
                score=1.0,
            ),
            # Fix 1: SEBI Registration Numbers (e.g. INM000013004, INR000004058)
            Pattern(
                name="sebi_registration_number",
                regex=r"\bIN[A-Z]\d{9,12}\b",
                score=1.0,
            ),
        ]
        super().__init__(
            supported_entity="SSN_GOVT_ID",
            patterns=patterns,
            name="CustomGovtIDRecognizer",
        )


class CustomDateOfBirthRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="full_month_date",
                regex=r"(?i)\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b",
                score=0.85,
            ),
            Pattern(
                name="short_month_date",
                regex=r"(?i)\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}\b",
                score=0.85,
            ),
            Pattern(
                name="numeric_date",
                regex=r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
                score=0.85,
            ),
            Pattern(
                name="iso_date",
                regex=r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b",
                score=0.85,
            ),
        ]
        super().__init__(
            supported_entity="DATE_OF_BIRTH",
            patterns=patterns,
            name="CustomDateOfBirthRecognizer",
        )


class NamePatternRecognizer(PatternRecognizer):
    def __init__(self, **kwargs):
        patterns = [
            Pattern(
                name="title_case_name_sequence",
                regex=r"\b[A-Z][a-zA-Z]{1,}(?:\s+[A-Z][a-zA-Z]{1,}){1,3}\b",
                score=0.85,
            ),
            Pattern(
                name="all_caps_name_sequence",
                regex=r"\b[A-Z]{2,}(?:\s+[A-Z]{2,}){1,3}\b",
                score=0.85,
            ),
        ]
        super().__init__(
            supported_entity="PERSON",
            patterns=patterns,
            name="NamePatternRecognizer",
            context=[
                "OUR PROMOTERS:",
                "PROMOTERS:",
                "PROMOTER:",
                "CONTACT PERSON:",
                "NAME OF THE",
                "SHAREHOLDER:",
                "SHAREHOLDERS:",
                "DIRECTOR:",
                "DIRECTORS:",
                "SECRETARY:",
                "COMPLIANCE OFFICER",
            ],
        )

    def validate_result(self, pattern_text: str) -> bool:
        words = [w.strip() for w in pattern_text.split() if w.strip()]
        if not (2 <= len(words) <= 4):
            return False

        if not all(w[0].isupper() for w in words):
            return False

        word_lowers = [w.lower() for w in words]
        word_uppers = [w.upper() for w in words]

        # 1. Reject if ANY word is a preposition, conjunction, or article
        preps_conjs = {
            "AND", "OR", "OF", "TO", "FOR", "IN", "ON", "AT", "BY", "WITH",
            "FROM", "INTO", "THE", "A", "AN", "AMONG", "ABOUT", "UNDER", "OVER"
        }
        if any(w in preps_conjs for w in word_uppers):
            return False

        # 2. Reject corporate suffixes
        corp_suffixes = {"LIMITED", "TRUST", "INC", "CORPORATION", "LLP", "PRIVATE", "HOLDINGS", "GROUP", "BANK", "PLC", "PVT", "LTD"}
        if any(w in corp_suffixes for w in word_uppers):
            return False

        # 3. Reject if ANY word is in NON_PII_WORDS or ALL_CAPS_STOP_WORDS
        if any(w in NON_PII_WORDS for w in word_lowers) or any(w in ALL_CAPS_STOP_WORDS for w in word_uppers):
            return False

        # 4. Reject common non-name terms (section titles, geographic terms, financial terms)
        non_name_words = {
            "government", "central", "state", "taluka", "district", "village", "city",
            "town", "road", "street", "building", "floor", "park", "complex",
            "section", "table", "details", "offer", "sale", "issue", "public",
            "acquisition", "cost", "weighted", "average", "statements", "looking",
            "forward", "market", "currency", "presentation", "financial", "position",
            "results", "operations", "summary", "general", "capital", "structure",
            "business", "objects", "basis", "price", "tax", "benefits", "legal",
            "information", "regulatory", "disclosures", "terms", "restrictions",
            "ownership", "foreign", "indian", "securities", "exchange", "board",
            "overview", "industry", "provisions", "main", "articles", "association",
        }
        if any(w in non_name_words for w in word_lowers):
            return False

        return True


# Fix 2: Custom Organization Pattern Recognizer for Indian Corporate Names
class CustomOrganizationRecognizer(PatternRecognizer):
    """Recognizes Indian corporate entity names by their legal suffixes (Private Limited, LLP, etc.)."""
    def __init__(self, **kwargs):
        patterns = [
            # "X Y Z Private Limited" or "X Y Z Pvt Ltd" or "X Y Z Pvt. Ltd."
            Pattern(
                name="indian_private_limited",
                regex=r"\b(?:[A-Z][A-Za-z&]+(?:\s+[A-Za-z&]+){0,8})\s+(?:Private\s+Limited|Pvt\.?\s*Ltd\.?)\b",
                score=0.95,
            ),
            # "X Y Z Limited" (public company)
            Pattern(
                name="indian_public_limited",
                regex=r"\b(?:[A-Z][A-Za-z&]+(?:\s+[A-Za-z&]+){0,8})\s+Limited\b",
                score=0.90,
            ),
            # "X & Y LLP" or "X Y LLP"
            Pattern(
                name="llp_entity",
                regex=r"\b(?:[A-Z][A-Za-z&]+(?:\s+[A-Za-z&]+){0,6})\s+LLP\b",
                score=0.95,
            ),
        ]
        super().__init__(
            supported_entity="ORGANIZATION",
            patterns=patterns,
            name="CustomOrganizationRecognizer",
        )

    def validate_result(self, pattern_text: str) -> bool:
        """Filter out matches that are clearly not organization names."""
        clean = pattern_text.strip()
        words = clean.split()
        if len(words) < 2:
            return False
        # Reject if first word is a common article/pronoun/preposition (not a proper noun)
        non_name_starters = {
            "the", "a", "an", "this", "that", "its", "our", "their", "his", "her",
            "and", "or", "but", "for", "with", "from", "into", "by", "as", "if",
            "is", "was", "are", "were", "be", "been", "being", "has", "have", "had",
            "it", "we", "they", "he", "she", "you", "who", "which", "what",
        }
        if words[0].lower() in non_name_starters:
            return False
        # Reject if match contains common verbs (indicates a sentence, not a name)
        verb_indicators = {
            "converted", "changed", "registered", "incorporated", "located",
            "issued", "filed", "established", "formed", "appointed", "resigned",
        }
        if any(w.lower() in verb_indicators for w in words):
            return False
        return True


# Fix 3: Custom Location Recognizer with Indian City Gazetteer
class CustomLocationRecognizer(PatternRecognizer):
    """Recognizes Indian location references: area/locality + known city name patterns."""

    INDIAN_CITIES = {
        "Mumbai", "Pune", "Delhi", "Bangalore", "Bengaluru", "Chennai",
        "Hyderabad", "Kolkata", "Ahmedabad", "Jaipur", "Lucknow",
        "Chandigarh", "Noida", "Gurgaon", "Gurugram", "Thane",
        "Navi Mumbai", "Surat", "Indore", "Nagpur", "Vadodara",
        "Kochi", "Coimbatore", "Visakhapatnam", "Bhopal", "Patna",
    }

    def __init__(self, **kwargs):
        # Build city alternation pattern
        city_alts = "|".join(sorted(self.INDIAN_CITIES, key=len, reverse=True))
        patterns = [
            # "AreaName CityName" — e.g. "Birdewadi Pune", "Baner Pune"
            Pattern(
                name="area_city_pattern",
                regex=rf"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){{0,3}}\s+(?:{city_alts})\b",
                score=0.92,
            ),
        ]
        super().__init__(
            supported_entity="LOCATION",
            patterns=patterns,
            name="CustomLocationRecognizer",
        )

    def validate_result(self, pattern_text: str) -> bool:
        """Ensure match ends with a known Indian city and is not a person name."""
        clean = pattern_text.strip()
        words = clean.split()
        if len(words) < 2:
            return False
        # Last word(s) must be a known city
        last_word = words[-1]
        if last_word not in self.INDIAN_CITIES:
            # Check for 2-word cities like "Navi Mumbai"
            if len(words) >= 3:
                two_word_city = f"{words[-2]} {words[-1]}"
                if two_word_city not in self.INDIAN_CITIES:
                    return False
            else:
                return False
        return True


# ---------------------------------------------------------------------------
# Standardized Predictable Pseudonym Mapper
# ---------------------------------------------------------------------------
class PseudonymMapper:
    """
    Thread-safe, execution-scoped persistent pseudonym dictionary.
    Uses clean, predictable, standardized placeholder pools mapped sequentially.
    """

    POOLS = {
        "PERSON": ["John Doe", "Jane Doe", "Peter Parker"],
        "FULL_NAMES": ["John Doe", "Jane Doe", "Peter Parker"],
        "EMAIL": ["john.doe@example.com", "jane.doe@example.com"],
        "EMAIL_ADDRESS": ["john.doe@example.com", "jane.doe@example.com"],
        "PHONE": ["+91 1234567890", "+91 9876543210"],
        "PHONE_NUMBER": ["+91 1234567890", "+91 9876543210"],
        "LOCATION": ["123, Confidential Street, City"],
        "ADDRESSES": ["123, Confidential Street, City"],
        "ORGANIZATION": ["Acme Corp", "Global Industries Ltd"],
        "COMPANY_NAMES": ["Acme Corp", "Global Industries Ltd"],
        "SSN_GOVT_ID": ["[REDACTED_ID]"],
        "GOVT_ID": ["[REDACTED_ID]"],
        "DATE_OF_BIRTH": ["01/01/1990"],
        "DATE_TIME": ["01/01/1990"],
        "IP_ADDRESS": ["192.168.1.1"],
        "CREDIT_CARD": ["4111-XXXX-XXXX-1111"],
    }

    def __init__(self, locale: str = "en_US", seed: Optional[int] = None):
        self._lock = threading.Lock()
        self._mapping: Dict[Tuple[str, str], str] = {}
        self._category_indices: Dict[str, int] = {}

    def get_pseudonym(self, original_text: str, entity_type: str) -> str:
        clean_text = original_text.strip()
        et = entity_type.upper()
        key = (et, clean_text)

        with self._lock:
            if key in self._mapping:
                return self._mapping[key]

            pool = self.POOLS.get(et) or self.POOLS.get(entity_type)
            if pool:
                idx = self._category_indices.get(et, 0)
                fake_val = pool[idx % len(pool)]
                self._category_indices[et] = idx + 1
            else:
                fake_val = f"[REDACTED_{et}]"

            self._mapping[key] = fake_val
            return fake_val

    def get_all_mappings(self) -> Dict[str, str]:
        with self._lock:
            return {orig: fake for (_, orig), fake in self._mapping.items()}


# ---------------------------------------------------------------------------
# Core PIIRedactor Engine
# ---------------------------------------------------------------------------
@dataclass
class DetectedEntity:
    entity_type: str
    original_text: str
    pseudonym: str
    confidence_score: float
    start: int
    end: int


class PIIRedactor:
    """
    Core PII Engine combining Presidio Analyzer, spaCy NER, custom pattern recognizers,
    contextual score boosting, and python-docx layout preservation.
    """

    CONTEXT_TRIGGERS = [
        "Promoter:",
        "Promoter",
        "Compliance Officer",
        "Registered Office",
        "Pvt Ltd",
        "Limited",
        "LLP",
        "Director:",
        "Authorized Signatory",
        "CIN:",
    ]

    # Fix 7: Include custom LOCATION and ORGANIZATION recognizers as high-priority
    # so they win overlap resolution against generic spaCy PERSON tags
    DETERMINISTIC_ENTITIES = {
        "EMAIL",
        "PHONE",
        "IP_ADDRESS",
        "CREDIT_CARD",
        "SSN_GOVT_ID",
        "DATE_OF_BIRTH",
        "LOCATION",
    }

    def __init__(self, spacy_model: str = "en_core_web_sm"):
        # Setup spaCy NLP provider for Presidio
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": spacy_model}],
        }
        provider = NlpEngineProvider(nlp_configuration=configuration)
        nlp_engine = provider.create_engine()

        # Initialize Recognizer Registry
        registry = RecognizerRegistry()
        registry.load_predefined_recognizers(nlp_engine=nlp_engine)

        # Register custom recognizers
        registry.add_recognizer(CustomEmailRecognizer())
        registry.add_recognizer(CustomPhoneRecognizer())
        registry.add_recognizer(CustomIPAddressRecognizer())
        registry.add_recognizer(CustomCreditCardRecognizer())
        registry.add_recognizer(CustomGovtIDRecognizer())
        registry.add_recognizer(CustomDateOfBirthRecognizer())
        registry.add_recognizer(NamePatternRecognizer())
        # Fix 2 & 3: Register custom organization and location recognizers
        registry.add_recognizer(CustomOrganizationRecognizer())
        registry.add_recognizer(CustomLocationRecognizer())

        self.analyzer = AnalyzerEngine(
            nlp_engine=nlp_engine,
            registry=registry,
            supported_languages=["en"]
        )
        self.nlp = nlp_engine.nlp.get("en") if hasattr(nlp_engine, "nlp") else None

        # Cache for paragraph text analysis to eliminate redundant passes
        self._analysis_cache: Dict[str, List[RecognizerResult]] = {}
        self._cache_lock = threading.Lock()

        # Map Presidio standard entities to requirement names
        self.entity_type_map = {
            "PERSON": "FULL_NAMES",
            "ORGANIZATION": "COMPANY_NAMES",
            "LOCATION": "ADDRESSES",
            "EMAIL_ADDRESS": "EMAIL",
            "PHONE_NUMBER": "PHONE",
        }

    def analyze_text(self, text: str, is_name_col: bool = False) -> List[RecognizerResult]:
        """Analyze text with Presidio using score_threshold=0.80, fast pre-filtering, promoter extractor, and LRU caching."""
        if not text or not text.strip():
            return []

        clean_t = text.strip()
        # Fast pre-filtering: skip short strings (<3 chars) or pure numbers/symbols without candidate PII
        if len(clean_t) < 3:
            return []

        has_letters = any(c.isalpha() for c in clean_t)
        has_pii_symbols = "@" in clean_t or "+" in clean_t or any(c.isdigit() for c in clean_t)
        if not (has_letters or has_pii_symbols):
            return []

        cache_key = (text, is_name_col)
        # Check analysis cache
        with self._cache_lock:
            if cache_key in self._analysis_cache:
                return self._analysis_cache[cache_key]

        # Analyze with Presidio using strict score_threshold=0.80 for NER entities
        results = self.analyzer.analyze(
            text=text,
            language="en",
            score_threshold=0.80,
        )

        # 1. Contextual ALL-CAPS & Title-Cased Promoters List Extractor
        promoter_match = re.search(r'\b(?:OUR\s+)?PROMOTERS\s*:\s*', text, re.IGNORECASE)
        if promoter_match:
            start_pos = promoter_match.end()
            promoter_section = text[start_pos:]
            segments = re.split(r',|\bAND\b|\band\b', promoter_section)
            search_offset = start_pos

            for seg in segments:
                seg_clean = seg.strip()
                if not seg_clean:
                    continue
                cleaned_name = re.sub(r'^[^\w]+|[^\w]+$', '', seg_clean)
                words = [w.strip() for w in cleaned_name.split() if w.strip()]

                suffix_words = {"FAMILY", "TRUST", "LIMITED", "PRIVATE", "CORPORATION", "LLP", "INC", "PLC", "HOLDINGS", "GROUP"}
                has_corp_suffix = any(w.upper() in suffix_words for w in words)

                if 2 <= len(words) <= 4 and not has_corp_suffix:
                    if all(w.isupper() or w.istitle() for w in words):
                        if not any(w.lower() in NON_PII_WORDS for w in words):
                            seg_match = re.search(re.escape(cleaned_name), text[search_offset:])
                            if seg_match:
                                abs_start = search_offset + seg_match.start()
                                abs_end = search_offset + seg_match.end()
                                search_offset = abs_end
                                results.append(
                                    RecognizerResult(
                                        entity_type="PERSON",
                                        start=abs_start,
                                        end=abs_end,
                                        score=0.98,
                                    )
                                )

        # 2. Table Cell Name Recognition (if is_name_col is True, redact cell text whatsoever)
        if is_name_col and clean_t:
            start_idx = text.find(clean_t)
            if start_idx != -1:
                results.append(
                    RecognizerResult(
                        entity_type="PERSON",
                        start=start_idx,
                        end=start_idx + len(clean_t),
                        score=1.0,
                    )
                )

        # Contextual Score Boosting for Legal/Corporate Triggers
        boosted_results: List[RecognizerResult] = []
        has_trigger = any(trigger in text for trigger in self.CONTEXT_TRIGGERS)

        for res in results:
            score = res.score
            if has_trigger and res.entity_type in ("PERSON", "ORGANIZATION", "LOCATION"):
                score = min(1.0, score + 0.20)
                res.score = score
            boosted_results.append(res)

        # Filter out HEADING_BLOCKLIST, NON_PII_WORDS, NON_PII_TERMS, monetary amounts, section headers, and apply Single-Word Filter
        filtered_results: List[RecognizerResult] = []
        for res in boosted_results:
            raw_match = text[res.start:res.end].strip()
            match_lower = raw_match.lower()
            words = [w.strip() for w in re.split(r'\s+', raw_match) if w.strip()]
            word_lowers = [w.lower() for w in words]

            # 1. Heading Blocklist Check
            if match_lower in HEADING_BLOCKLIST or any(heading == match_lower or match_lower in heading for heading in HEADING_BLOCKLIST):
                continue

            # 2. Strict Confidence Thresholds
            if res.entity_type in ("PERSON", "ORGANIZATION", "LOCATION"):
                if res.score < 0.80:
                    continue

            # Fix 5: Filter short abbreviation ORG detections (e.g. "CIN", "BSE")
            if res.entity_type == "ORGANIZATION" and len(raw_match) <= 3:
                continue

            # 3. Single-Word Filter for PERSON (require explicit honorific prefix unless inside a name column)
            if res.entity_type == "PERSON" and not is_name_col:
                if len(words) == 1:
                    prefix_ctx = text[max(0, res.start - 20):res.start].strip().lower()
                    honorifics = ["mr.", "mr", "ms.", "ms", "mrs.", "mrs", "dr.", "dr", "smt.", "smt", "shri", "prof.", "prof"]
                    has_honorific = any(prefix_ctx.endswith(h) or (prefix_ctx.split() and prefix_ctx.split()[-1] == h) for h in honorifics)
                    if not has_honorific:
                        continue

            # 4. Non-PII Boilerplate Exclusion Filter
            # Fix 6: For ORGANIZATION, only reject if ALL words are stop words (not partial matches)
            #         This prevents filtering out valid orgs like "ICICI Securities Limited"
            if res.entity_type in ("PERSON", "LOCATION"):
                if match_lower in NON_PII_WORDS or match_lower in NON_PII_TERMS:
                    continue
                if len(words) == 1 and word_lowers[0] in NON_PII_WORDS:
                    continue
                if all(w in NON_PII_WORDS or w.upper() in ALL_CAPS_STOP_WORDS for w in word_lowers):
                    continue
            elif res.entity_type == "ORGANIZATION":
                if match_lower in NON_PII_TERMS:
                    continue
                # Only reject orgs where ALL words are generic stop words
                if all(w in NON_PII_WORDS or w.upper() in ALL_CAPS_STOP_WORDS for w in word_lowers):
                    continue

            # 5. Contextual Exclusions (monetary, page numbers, statutory citations)
            prefix_ctx = text[max(0, res.start - 25):res.start].lower()
            suffix_ctx = text[res.end:min(len(text), res.end + 25)].lower()

            if any(term in prefix_ctx or term in suffix_ctx for term in ["₹", "rs.", "rupees", "million", "crore", "lakh", "section", "clause", "page"]):
                if res.entity_type in ("PHONE", "DATE_OF_BIRTH", "LOCATION"):
                    continue

            # Protect Statutory Citation Years (e.g. "Companies Act, 2013", "SEBI Regulations, 2018")
            if res.entity_type in ("DATE_OF_BIRTH", "DATE_TIME") and re.match(r"^\d{4}$", raw_match):
                if any(stat in prefix_ctx for stat in ["act", "regulations", "rules", "section", "clause", "dated", "year", "fy", "prospectus"]):
                    continue

            filtered_results.append(res)

        # Fix 8: Store in cache using consistent cache_key tuple (text, is_name_col)
        with self._cache_lock:
            if len(self._analysis_cache) > 5000:
                self._analysis_cache.clear()
            self._analysis_cache[cache_key] = filtered_results

        return filtered_results

    def _resolve_overlapping_entities(
        self, results: List[RecognizerResult]
    ) -> List[RecognizerResult]:
        """
        Resolve overlapping entity intervals. Deterministic regex rules take precedence
        over generic spaCy NER tags when spans overlap.
        """
        if not results:
            return []

        # Helper to determine priority
        def priority_tuple(r: RecognizerResult):
            is_deterministic = r.entity_type in self.DETERMINISTIC_ENTITIES
            span_len = r.end - r.start
            return (r.start, 1 if is_deterministic else 0, r.score, span_len)

        # Sort by start position ascending, deterministic preference, then score
        sorted_results = sorted(results, key=priority_tuple)

        resolved: List[RecognizerResult] = []
        for res in sorted_results:
            if not resolved:
                resolved.append(res)
                continue

            last = resolved[-1]
            # Check overlap: max(start) < min(end)
            if max(res.start, last.start) < min(res.end, last.end):
                last_is_det = last.entity_type in self.DETERMINISTIC_ENTITIES
                res_is_det = res.entity_type in self.DETERMINISTIC_ENTITIES

                if res_is_det and not last_is_det:
                    resolved[-1] = res
                elif last_is_det and not res_is_det:
                    pass
                else:
                    if res.score > last.score:
                        resolved[-1] = res
            else:
                resolved.append(res)

        return resolved

    def redact_paragraph(
        self, paragraph: Paragraph, pseudonym_mapper: PseudonymMapper, is_name_col: bool = False
    ) -> List[DetectedEntity]:
        """
        Redact PII in a docx Paragraph while preserving run-level formatting
        (bold, italics, font size, font family, colors).
        """
        full_text = paragraph.text
        if not full_text or not full_text.strip():
            return []

        analysis_results = self.analyze_text(full_text, is_name_col=is_name_col)
        resolved_results = self._resolve_overlapping_entities(analysis_results)

        if not resolved_results:
            return []

        detected_entities: List[DetectedEntity] = []
        replacements: List[Tuple[int, int, str]] = []

        for res in resolved_results:
            orig_text = full_text[res.start:res.end]
            mapped_type = self.entity_type_map.get(res.entity_type, res.entity_type)
            fake_val = pseudonym_mapper.get_pseudonym(orig_text, mapped_type)

            detected_entities.append(
                DetectedEntity(
                    entity_type=mapped_type,
                    original_text=orig_text,
                    pseudonym=fake_val,
                    confidence_score=res.score,
                    start=res.start,
                    end=res.end,
                )
            )
            replacements.append((res.start, res.end, fake_val))

        # Perform run-level substitution from right to left (end to start)
        self._apply_run_replacements(paragraph, replacements)
        return detected_entities

    def _apply_run_replacements(
        self, paragraph: Paragraph, replacements: List[Tuple[int, int, str]]
    ) -> None:
        """
        Safely replace text intervals in a paragraph across runs.
        Processes replacements in reverse order of start index to maintain valid offsets.
        """
        if not paragraph.runs or not replacements:
            return

        # Sort replacements by start index descending
        sorted_replacements = sorted(replacements, key=lambda r: r[0], reverse=True)

        for start, end, replacement in sorted_replacements:
            # Build current run bounds
            run_bounds: List[Tuple[int, int]] = []
            curr_pos = 0
            for run in paragraph.runs:
                run_len = len(run.text)
                run_bounds.append((curr_pos, curr_pos + run_len))
                curr_pos += run_len

            # Find target runs intersecting [start, end)
            target_run_indices = [
                i
                for i, (r_start, r_end) in enumerate(run_bounds)
                if max(start, r_start) < min(end, r_end)
            ]

            if not target_run_indices:
                continue

            first_idx = target_run_indices[0]
            last_idx = target_run_indices[-1]

            first_run = paragraph.runs[first_idx]
            first_r_start, _ = run_bounds[first_idx]
            prefix = first_run.text[: start - first_r_start]

            if first_idx == last_idx:
                suffix = first_run.text[end - first_r_start :]
                first_run.text = prefix + replacement + suffix
            else:
                last_run = paragraph.runs[last_idx]
                last_r_start, _ = run_bounds[last_idx]
                suffix = last_run.text[end - last_r_start :]

                first_run.text = prefix + replacement
                last_run.text = suffix

                # Clear text in intermediate runs
                for mid_idx in range(first_idx + 1, last_idx):
                    paragraph.runs[mid_idx].text = ""

    def process_table(
        self, table: Any, pseudonym_mapper: PseudonymMapper, visited_cells: Optional[set] = None
    ) -> List[DetectedEntity]:
        """Safely iterate through docx table cells and cell paragraphs, avoiding duplicate merged cells."""
        if visited_cells is None:
            visited_cells = set()

        HEADER_NAME_KEYWORDS = [
            "NAME", "PROMOTER", "PROMOTOR", "CONTACT", "DIRECTOR", "OFFICER",
            "SHAREHOLDER", "EXECUTIVE", "MANAGEMENT", "KMP", "SIGNATORY", "AUDITOR"
        ]

        # Identify column indices where any top 3 row header cell contains a name/contact keyword
        name_col_indices = set()
        header_row_count = 1

        if table.rows:
            scan_rows = min(3, len(table.rows))
            for r_idx in range(scan_rows):
                row_has_keyword = False
                for col_idx, cell in enumerate(table.rows[r_idx].cells):
                    hdr_txt = cell.text.strip().upper()
                    if any(k in hdr_txt for k in HEADER_NAME_KEYWORDS):
                        name_col_indices.add(col_idx)
                        row_has_keyword = True
                if row_has_keyword:
                    header_row_count = max(header_row_count, r_idx + 1)

        entities: List[DetectedEntity] = []
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_id = id(cell._element)
                if cell_id in visited_cells:
                    continue
                visited_cells.add(cell_id)

                is_name_column = (c_idx in name_col_indices) and (r_idx >= header_row_count)
                for p in cell.paragraphs:
                    entities.extend(self.redact_paragraph(p, pseudonym_mapper, is_name_col=is_name_column))
                for nested_table in cell.tables:
                    entities.extend(self.process_table(nested_table, pseudonym_mapper, visited_cells))
        return entities

    def redact_docx(
        self,
        doc_input: Union[str, BinaryIO, bytes],
        pseudonym_mapper: Optional[PseudonymMapper] = None,
    ) -> Tuple[bytes, Dict[str, Any]]:
        """
        Redact a .docx document provided as a file path, stream, or bytes.

        Returns:
            Tuple[bytes, Dict[str, Any]]: Redacted document bytes and JSON metadata summary.
        """
        if pseudonym_mapper is None:
            pseudonym_mapper = PseudonymMapper()

        if isinstance(doc_input, bytes):
            doc_stream = io.BytesIO(doc_input)
        elif isinstance(doc_input, str):
            with open(doc_input, "rb") as f:
                doc_stream = io.BytesIO(f.read())
        else:
            doc_stream = doc_input

        doc = docx.Document(doc_stream)
        all_detected_entities: List[DetectedEntity] = []

        # Clear per-document analysis cache to maintain freshness
        with self._cache_lock:
            self._analysis_cache.clear()

        visited_cells = set()

        # 1. Main Document Paragraphs
        for p in doc.paragraphs:
            all_detected_entities.extend(self.redact_paragraph(p, pseudonym_mapper))

        # 2. Main Document Tables
        for table in doc.tables:
            all_detected_entities.extend(self.process_table(table, pseudonym_mapper, visited_cells))

        # 3. Document Sections (Headers and Footers)
        for section in doc.sections:
            for header in (
                section.header,
                getattr(section, "first_page_header", None),
                getattr(section, "even_page_header", None),
            ):
                if header is not None and not getattr(header, "is_linked_to_previous", False):
                    for hp in header.paragraphs:
                        all_detected_entities.extend(
                            self.redact_paragraph(hp, pseudonym_mapper)
                        )
                    for ht in header.tables:
                        all_detected_entities.extend(
                            self.process_table(ht, pseudonym_mapper, visited_cells)
                        )

            for footer in (
                section.footer,
                getattr(section, "first_page_footer", None),
                getattr(section, "even_page_footer", None),
            ):
                if footer is not None and not getattr(footer, "is_linked_to_previous", False):
                    for fp in footer.paragraphs:
                        all_detected_entities.extend(
                            self.redact_paragraph(fp, pseudonym_mapper)
                        )
                    for ft in footer.tables:
                        all_detected_entities.extend(
                            self.process_table(ft, pseudonym_mapper, visited_cells)
                        )

        # Save redacted document to BytesIO stream
        out_stream = io.BytesIO()
        doc.save(out_stream)
        redacted_bytes = out_stream.getvalue()

        # Build Metadata Summary
        entity_counts: Dict[str, int] = {}
        entity_occurrences: Dict[Tuple[str, str, str], int] = {}
        entity_scores: Dict[Tuple[str, str], float] = {}

        for ent in all_detected_entities:
            entity_counts[ent.entity_type] = entity_counts.get(ent.entity_type, 0) + 1
            key = (ent.entity_type, ent.original_text, ent.pseudonym)
            entity_occurrences[key] = entity_occurrences.get(key, 0) + 1
            entity_scores[(ent.entity_type, ent.original_text)] = ent.confidence_score

        detected_entities_list = [
            {
                "entity_type": ent_type,
                "original_text": orig,
                "pseudonym": fake,
                "confidence_score": entity_scores.get((ent_type, orig), 1.0),
                "count": count,
            }
            for (ent_type, orig, fake), count in entity_occurrences.items()
        ]

        metadata = {
            "total_entities_found": len(all_detected_entities),
            "entity_counts": entity_counts,
            "pseudonym_mappings": pseudonym_mapper.get_all_mappings(),
            "detected_entities": detected_entities_list,
        }

        return redacted_bytes, metadata


# ---------------------------------------------------------------------------
# Singleton Instance & Module Entrypoint
# ---------------------------------------------------------------------------
_REDACTOR_INSTANCE: Optional[PIIRedactor] = None
_REDACTOR_LOCK = threading.Lock()


def get_redactor(spacy_model: str = "en_core_web_sm") -> PIIRedactor:
    """Thread-safe singleton getter for PIIRedactor instance."""
    global _REDACTOR_INSTANCE
    if _REDACTOR_INSTANCE is None:
        with _REDACTOR_LOCK:
            if _REDACTOR_INSTANCE is None:
                _REDACTOR_INSTANCE = PIIRedactor(spacy_model=spacy_model)
    return _REDACTOR_INSTANCE


def redact_document(
    doc_input: Union[str, BinaryIO, bytes],
    spacy_model: str = "en_core_web_sm",
    pseudonym_mapper: Optional[PseudonymMapper] = None,
) -> Tuple[bytes, Dict[str, Any]]:
    """
    Convenience function to redact a docx document using pre-warmed singleton engine.
    """
    engine = get_redactor(spacy_model=spacy_model)
    return engine.redact_docx(doc_input, pseudonym_mapper=pseudonym_mapper)
