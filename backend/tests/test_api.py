"""
API Integration Tests for FastAPI Backend Server (backend/app/main.py)
"""

import io
import os
import sys
import pytest
import docx
from fastapi.testclient import TestClient

# Ensure backend root is in sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from app.main import app

client = TestClient(app)


def test_root_endpoint():
    response = client.get("/")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "online"
    assert data["app"] == "PII Redaction Engine API"


def test_health_endpoint():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "healthy"}


def test_evaluate_endpoint():
    response = client.get("/evaluate")
    assert response.status_code == 200
    data = response.json()
    assert "overall" in data
    assert "precision" in data["overall"]
    assert "recall" in data["overall"]
    assert "f1_score" in data["overall"]
    assert "by_entity_type" in data


def test_redact_endpoint_binary():
    # Build sample DOCX in memory
    doc = docx.Document()
    doc.add_paragraph("Contact Kushal Subbayya Hegde at cs.connect@kshinternational.com or +91 98765 43210.")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    response = client.post(
        "/redact",
        files={"file": ("sample_test.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    assert "Content-Disposition" in response.headers
    assert 'filename="redacted_sample_test.docx"' in response.headers["Content-Disposition"]
    
    # Reload returned bytes to ensure it's a valid docx
    redacted_doc = docx.Document(io.BytesIO(response.content))
    assert len(redacted_doc.paragraphs) > 0


def test_redact_endpoint_metadata():
    # Build sample DOCX in memory
    doc = docx.Document()
    doc.add_paragraph("Director: Sarthak Malvadkar holds PAN ABCDE1234F.")
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    response = client.post(
        "/redact?return_metadata=true",
        files={"file": ("sample_meta.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    data = response.json()
    assert data["filename"] == "sample_meta.docx"
    assert data["redacted_filename"] == "redacted_sample_meta.docx"
    assert "redacted_file_base64" in data
    assert "metadata" in data
    assert data["metadata"]["total_entities_found"] > 0


def test_redact_invalid_file_extension():
    response = client.post(
        "/redact",
        files={"file": ("invalid_text.txt", b"Hello world", "text/plain")},
    )
    assert response.status_code == 400
    assert "Only Microsoft Word (.docx) documents are supported" in response.json()["detail"]
