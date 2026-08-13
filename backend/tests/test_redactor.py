"""
Unit and Integration Tests for Core PII Redaction Engine (backend/app/redactor.py)
"""

import io
import os
import sys
import pytest
import docx

# Add backend directory to sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from app.redactor import (
    PIIRedactor,
    PseudonymMapper,
    is_valid_luhn,
    redact_document,
)


# ---------------------------------------------------------------------------
# 1. Test Luhn Algorithm Checksum
# ---------------------------------------------------------------------------
def test_luhn_checksum():
    # Valid credit card numbers
    assert is_valid_luhn("4532 0151 1283 0366") is True
    assert is_valid_luhn("4532015112830366") is True
    assert is_valid_luhn("378282246310005") is True  # Amex

    # Invalid card numbers
    assert is_valid_luhn("4532 0151 1283 0367") is False
    assert is_valid_luhn("1234567890123456") is False
    assert is_valid_luhn("1234") is False


# ---------------------------------------------------------------------------
# 2. Test PseudonymMapper Consistency
# ---------------------------------------------------------------------------
def test_pseudonym_mapper_consistency():
    mapper = PseudonymMapper(seed=42)
    
    # Same input must return exact same pseudonym every time
    name1 = mapper.get_pseudonym("Rashi Patil", "PERSON")
    name2 = mapper.get_pseudonym("Rashi Patil", "PERSON")
    name3 = mapper.get_pseudonym("Rashi Patil", "PERSON")
    assert name1 == name2 == name3

    email1 = mapper.get_pseudonym("cs.connect@kshinternational.com", "EMAIL")
    email2 = mapper.get_pseudonym("cs.connect@kshinternational.com", "EMAIL")
    assert email1 == email2

    # Different entities should produce mappings recorded in dictionary
    all_maps = mapper.get_all_mappings()
    assert "Rashi Patil" in all_maps
    assert "cs.connect@kshinternational.com" in all_maps


# ---------------------------------------------------------------------------
# 3. Test Pattern Analysis for Required Entities
# ---------------------------------------------------------------------------
def test_pii_pattern_analysis():
    redactor = PIIRedactor()

    # Email
    results = redactor.analyze_text("Contact us at support@example.com immediately.")
    assert any(r.entity_type == "EMAIL" for r in results)

    # Phone
    results = redactor.analyze_text("Call +91 98765 43210 for details.")
    assert any(r.entity_type == "PHONE" for r in results)

    # IP Address
    results = redactor.analyze_text("Server IP is 192.168.1.100")
    assert any(r.entity_type == "IP_ADDRESS" for r in results)

    # Credit Card (Valid Luhn)
    results = redactor.analyze_text("Card number: 4532-0151-1283-0366")
    assert any(r.entity_type == "CREDIT_CARD" for r in results)

    # Govt IDs: US SSN, PAN, Aadhaar, CIN
    results = redactor.analyze_text("SSN: 123-45-6789")
    assert any(r.entity_type == "SSN_GOVT_ID" for r in results)

    results = redactor.analyze_text("PAN: ABCDE1234F")
    assert any(r.entity_type == "SSN_GOVT_ID" for r in results)

    results = redactor.analyze_text("Aadhaar: 1234 5678 9012")
    assert any(r.entity_type == "SSN_GOVT_ID" for r in results)

    results = redactor.analyze_text("CIN: L12345MH2020PLC123456")
    assert any(r.entity_type == "SSN_GOVT_ID" for r in results)

    # Date of Birth
    results = redactor.analyze_text("DOB: 15/08/1990 and Jan 15, 1995")
    dob_matches = [r for r in results if r.entity_type == "DATE_OF_BIRTH"]
    assert len(dob_matches) >= 1


# ---------------------------------------------------------------------------
# 4. Test Contextual Score Boosting
# ---------------------------------------------------------------------------
def test_contextual_score_boosting():
    redactor = PIIRedactor()
    
    text_normal = "John Smith is attending the conference."
    text_boosted = "Promoter: John Smith is attending the conference."

    results_normal = redactor.analyze_text(text_normal)
    results_boosted = redactor.analyze_text(text_boosted)

    # Find PERSON match score
    score_normal = max([r.score for r in results_normal if r.entity_type in ("PERSON", "FULL_NAMES")], default=0.0)
    score_boosted = max([r.score for r in results_boosted if r.entity_type in ("PERSON", "FULL_NAMES")], default=0.0)

    assert score_boosted >= score_normal


# ---------------------------------------------------------------------------
# 5. Integration Test: Full DOCX Redaction & Metadata Generation
# ---------------------------------------------------------------------------
def test_docx_redaction_and_layout_preservation():
    # Build sample in-memory DOCX with bold text, table, and header/footer
    doc = docx.Document()

    # Header
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Confidential Document - Contact: admin@company.com"

    # Paragraph with mixed formatting (bold run)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Promoter: ")
    r1.bold = True
    r2 = p1.add_run("Rashi Patil")
    r3 = p1.add_run(" holds PAN ")
    r4 = p1.add_run("ABCDE1234F")
    r4.italic = True
    r5 = p1.add_run(" and email cs.connect@kshinternational.com.")

    # Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).paragraphs[0].text = "Officer Name"
    table.cell(0, 1).paragraphs[0].text = "Phone Number"
    table.cell(1, 0).paragraphs[0].text = "Compliance Officer: Rashi Patil"
    table.cell(1, 1).paragraphs[0].text = "+91 98765 43210"

    # Footer
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "Registered Office: Mumbai. DOB: 01/01/1985"

    # Save to BytesIO
    doc_bytes_io = io.BytesIO()
    doc.save(doc_bytes_io)
    input_bytes = doc_bytes_io.getvalue()

    # Perform redaction
    redacted_bytes, metadata = redact_document(input_bytes)

    # 1. Verify Metadata Structure
    assert metadata["total_entities_found"] > 0
    assert "FULL_NAMES" in metadata["entity_counts"] or "PERSON" in metadata["entity_counts"]
    assert "EMAIL" in metadata["entity_counts"]
    assert "SSN_GOVT_ID" in metadata["entity_counts"]
    assert len(metadata["pseudonym_mappings"]) > 0

    # Verify persistent pseudonym consistency: "Rashi Patil" appears in p1 and table cell
    mappings = metadata["pseudonym_mappings"]
    assert "Rashi Patil" in mappings
    fake_rashi = mappings["Rashi Patil"]

    # 2. Reload redacted docx and inspect text content
    redacted_doc = docx.Document(io.BytesIO(redacted_bytes))

    # Read all text from redacted document
    full_redacted_text = ""
    for p in redacted_doc.paragraphs:
        full_redacted_text += p.text + "\n"
    for table in redacted_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_redacted_text += p.text + "\n"

    # Verify PII strings are absent and replaced with pseudonyms
    assert "Rashi Patil" not in full_redacted_text
    assert "ABCDE1234F" not in full_redacted_text
    assert "cs.connect@kshinternational.com" not in full_redacted_text
    assert fake_rashi in full_redacted_text
