#!/usr/bin/env python3
"""DOCX PII Redaction Script.

A document-agnostic, two-pass PII redaction pipeline:
- Pass 1: Scans the entire document to discover all entities, emails, domains, and promoter groups,
  dynamically building a synchronized session pseudonym registry without any hardcoded dictionary.
- Pass 2: Traverses all paragraphs, tables (with XML node deduplication for merged cells),
  headers, footers, and hyperlinks to perform formatting-preserving replacements.
"""

import os
import re
import sys
from typing import List, Tuple, Set
import docx
from pii_redactor import PIIRedactor


def replace_in_runs(runs, replacements: List[Tuple[str, str]]) -> bool:
    """Replaces text in a list of Runs while strictly preserving run formatting."""
    if not runs or not replacements:
        return False

    full_text = ''.join(r.text for r in runs)
    if not full_text:
        return False

    # Find all match positions in full text
    matches = []
    for orig, fake in replacements:
        start = 0
        while True:
            idx = full_text.find(orig, start)
            if idx == -1:
                break
            matches.append((idx, idx + len(orig), fake))
            start = idx + len(orig)

    if not matches:
        return False

    # Sort matches by start position and eliminate overlapping intervals
    matches.sort(key=lambda x: x[0])
    filtered_matches = []
    last_end = 0
    for m in matches:
        if m[0] >= last_end:
            filtered_matches.append(m)
            last_end = m[1]

    # Calculate run boundaries
    run_spans = []
    curr = 0
    for r in runs:
        r_len = len(r.text)
        run_spans.append((curr, curr + r_len))
        curr += r_len

    # Process matches from right to left so indices remain valid
    for m_start, m_end, fake in reversed(filtered_matches):
        first_run_idx = None
        last_run_idx = None
        for i, (r_start, r_end) in enumerate(run_spans):
            if r_end > m_start and r_start < m_end:
                if first_run_idx is None:
                    first_run_idx = i
                last_run_idx = i

        if first_run_idx is None:
            continue

        if first_run_idx == last_run_idx:
            r = runs[first_run_idx]
            r_start, _ = run_spans[first_run_idx]
            rel_start = m_start - r_start
            rel_end = m_end - r_start
            r.text = r.text[:rel_start] + fake + r.text[rel_end:]
        else:
            # First run: keep text before match + append fake replacement
            r1 = runs[first_run_idx]
            r1_start, _ = run_spans[first_run_idx]
            rel_start = m_start - r1_start
            r1.text = r1.text[:rel_start] + fake

            # Intermediate runs: clear text
            for i in range(first_run_idx + 1, last_run_idx):
                runs[i].text = ''

            # Last run: keep text after match
            r2 = runs[last_run_idx]
            r2_start, _ = run_spans[last_run_idx]
            rel_end = m_end - r2_start
            r2.text = r2.text[rel_end:]

        # Recompute run_spans for next iteration
        run_spans = []
        curr = 0
        for r in runs:
            r_len = len(r.text)
            run_spans.append((curr, curr + r_len))
            curr += r_len

    return True


def redact_paragraph(p, redactor: PIIRedactor) -> bool:
    """Redacts PII in a single paragraph while preserving run formatting."""
    text = p.text
    if not text or not text.strip():
        return False

    replacements = redactor.get_replacements(text)
    if not replacements:
        return False

    if p.runs:
        return replace_in_runs(p.runs, replacements)
    else:
        p.text = redactor.redact_text(text)
        return True


def redact_docx_file(input_path: str, output_path: str):
    """Reads input_path DOCX, discovers entities dynamically in Pass 1, and redacts in Pass 2."""
    print(f"Loading document: {input_path}")
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    doc = docx.Document(input_path)
    redactor = PIIRedactor()

    # Pass 1: Dynamic Discovery & Global Session Registry Population
    print("Pass 1: Dynamically scanning and discovering entities across document...")
    redactor.scan_and_mine_document(doc)
    print(f"Pass 1 Complete: Discovered and mapped {len(redactor.pseudonym_map)} entity variants.")

    # Pass 2: Redaction Execution
    print("Pass 2: Performing formatting-preserving redactions...")
    redacted_paragraph_count = 0
    seen_elements = set()

    # 1. Process Body Paragraphs
    print("Processing body paragraphs...")
    for p in doc.paragraphs:
        if p._p in seen_elements:
            continue
        seen_elements.add(p._p)
        if redact_paragraph(p, redactor):
            redacted_paragraph_count += 1

    # 2. Process Tables (Deduplicating merged cells using XML node identity)
    print("Processing tables (with merged cell deduplication)...")
    for table in doc.tables:
        seen_tc = set()
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_tc:
                    continue
                seen_tc.add(cell._tc)
                for p in cell.paragraphs:
                    if p._p in seen_elements:
                        continue
                    seen_elements.add(p._p)
                    if redact_paragraph(p, redactor):
                        redacted_paragraph_count += 1

    # 3. Process Section Headers & Footers
    print("Processing headers & footers...")
    for section in doc.sections:
        # Header
        if section.header:
            for p in section.header.paragraphs:
                if p._p not in seen_elements:
                    seen_elements.add(p._p)
                    if redact_paragraph(p, redactor):
                        redacted_paragraph_count += 1
            for table in section.header.tables:
                seen_tc = set()
                for row in table.rows:
                    for cell in row.cells:
                        if cell._tc in seen_tc:
                            continue
                        seen_tc.add(cell._tc)
                        for p in cell.paragraphs:
                            if p._p not in seen_elements:
                                seen_elements.add(p._p)
                                if redact_paragraph(p, redactor):
                                    redacted_paragraph_count += 1
        # Footer
        if section.footer:
            for p in section.footer.paragraphs:
                if p._p not in seen_elements:
                    seen_elements.add(p._p)
                    if redact_paragraph(p, redactor):
                        redacted_paragraph_count += 1
            for table in section.footer.tables:
                seen_tc = set()
                for row in table.rows:
                    for cell in row.cells:
                        if cell._tc in seen_tc:
                            continue
                        seen_tc.add(cell._tc)
                        for p in cell.paragraphs:
                            if p._p not in seen_elements:
                                seen_elements.add(p._p)
                                if redact_paragraph(p, redactor):
                                    redacted_paragraph_count += 1

    # 4. Redact Relationship Hyperlinks (w:hyperlink targets)
    print("Processing relationship hyperlinks...")
    try:
        for rel in doc.part.rels.values():
            if rel.is_external and rel.target_ref:
                target = rel.target_ref
                for domain, fake_domain in redactor.domains_map.items():
                    if domain in target:
                        rel._target = target.replace(domain, fake_domain)
                        break
    except Exception as e:
        print(f"Note on hyperlink relationship update: {e}")

    # Save output file
    doc.save(output_path)
    print(f"\nRedaction complete! Total modified elements: {redacted_paragraph_count}")
    print(f"Redacted document saved to: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        input_docx = "Red Herring Prospectus.docx"
        output_docx = "Output_Redacted_Prospectus.docx"
    else:
        input_docx = sys.argv[1]
        output_docx = sys.argv[2]

    redact_docx_file(input_docx, output_docx)